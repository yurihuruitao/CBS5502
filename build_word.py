"""
Combine METHODOLOGY_EN.md and RESULTS_EN.md into a single Word document with images.
"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

ROOT = Path(__file__).parent

def parse_md_blocks(text):
    """Parse markdown into a list of blocks: heading, paragraph, table, image, note."""
    lines = text.split('\n')
    blocks = []
    i = 0
    while i < len(lines):
        line = lines[i]

        # Heading
        m = re.match(r'^(#{1,4})\s+(.*)', line)
        if m:
            level = len(m.group(1))
            blocks.append(('heading', level, m.group(2).strip()))
            i += 1
            continue

        # Image
        m = re.match(r'^\s*!\[([^\]]*)\]\(([^)]+)\)', line)
        if m:
            blocks.append(('image', m.group(2), m.group(1)))
            i += 1
            continue

        # Table (detect by |)
        if line.strip().startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            blocks.append(('table', table_lines))
            continue

        # Horizontal rule
        if re.match(r'^-{3,}\s*$', line.strip()):
            i += 1
            continue

        # Blockquote
        if line.strip().startswith('>'):
            text_content = line.strip().lstrip('>').strip()
            i += 1
            while i < len(lines) and lines[i].strip().startswith('>'):
                text_content += ' ' + lines[i].strip().lstrip('>').strip()
                i += 1
            blocks.append(('blockquote', text_content))
            continue

        # Empty line
        if not line.strip():
            i += 1
            continue

        # Regular paragraph (collect consecutive non-empty lines)
        para_lines = []
        while i < len(lines) and lines[i].strip() and not lines[i].strip().startswith('#') \
                and not lines[i].strip().startswith('|') and not re.match(r'^\s*!\[', lines[i]) \
                and not lines[i].strip().startswith('>') and not re.match(r'^-{3,}\s*$', lines[i].strip()):
            para_lines.append(lines[i])
            i += 1
        if para_lines:
            blocks.append(('paragraph', ' '.join(para_lines)))

    return blocks


def add_rich_text(paragraph, text):
    """Add text with bold and italic markdown formatting to a paragraph."""
    # Process bold+italic (***text*** or ___text___), bold (**text**), italic (*text* or _text_)
    pattern = re.compile(
        r'\*\*\*(.+?)\*\*\*'   # ***bold italic***
        r'|\*\*(.+?)\*\*'      # **bold**
        r'|\*(.+?)\*'          # *italic*
        r'|`([^`]+)`'          # `code`
    )
    pos = 0
    for m in pattern.finditer(text):
        # Add text before match
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        if m.group(1):  # bold italic
            r = paragraph.add_run(m.group(1))
            r.bold = True
            r.italic = True
        elif m.group(2):  # bold
            r = paragraph.add_run(m.group(2))
            r.bold = True
        elif m.group(3):  # italic
            r = paragraph.add_run(m.group(3))
            r.italic = True
        elif m.group(4):  # code
            r = paragraph.add_run(m.group(4))
            r.font.name = 'Courier New'
            r.font.size = Pt(9)
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def parse_table(table_lines):
    """Parse markdown table lines into header and rows."""
    rows = []
    for line in table_lines:
        cells = [c.strip() for c in line.strip().strip('|').split('|')]
        rows.append(cells)
    # Remove separator row (contains ---)
    rows = [r for r in rows if not all(re.match(r'^[-:]+$', c) for c in r)]
    return rows


def add_table(doc, rows):
    """Add a formatted table to the document."""
    if not rows:
        return
    ncols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            cell = table.cell(i, j)
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Clean markdown formatting for table cells
            clean = cell_text.strip()
            if clean.startswith('**') and clean.endswith('**'):
                r = p.add_run(clean.strip('*'))
                r.bold = True
            else:
                add_rich_text(p, clean)
            # Font size
            for run in p.runs:
                run.font.size = Pt(9)

    # Bold header row
    for cell in table.rows[0].cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True


def build_doc():
    doc = Document()

    # Page setup
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # Default font
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    # Heading styles
    for i in range(1, 5):
        hs = doc.styles[f'Heading {i}']
        hs.font.name = 'Times New Roman'
        hs.font.color.rgb = None  # remove blue color
        if i == 1:
            hs.font.size = Pt(16)
        elif i == 2:
            hs.font.size = Pt(13)
        elif i == 3:
            hs.font.size = Pt(11.5)
        else:
            hs.font.size = Pt(11)

    # Process both files
    for md_file in ['METHODOLOGY_EN.md', 'RESULTS_EN.md']:
        text = (ROOT / md_file).read_text(encoding='utf-8')

        # Strip references section from methodology (will be added once at end)
        if md_file == 'METHODOLOGY_EN.md':
            ref_idx = text.find('\n## References')
            if ref_idx != -1:
                references_text = text[ref_idx:]
                text = text[:ref_idx]
            else:
                references_text = ''

        blocks = parse_md_blocks(text)

        for block in blocks:
            btype = block[0]

            if btype == 'heading':
                _, level, content = block
                p = doc.add_heading(content, level=level)
                # Remove any auto-coloring
                for run in p.runs:
                    run.font.color.rgb = None

            elif btype == 'paragraph':
                _, content = block
                p = doc.add_paragraph()
                add_rich_text(p, content)

            elif btype == 'blockquote':
                _, content = block
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(1.5)
                add_rich_text(p, content)

            elif btype == 'image':
                _, img_path, alt_text = block
                full_path = ROOT / img_path
                if full_path.exists():
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    run.add_picture(str(full_path), width=Inches(5.5))
                else:
                    p = doc.add_paragraph(f'[Image not found: {img_path}]')
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            elif btype == 'table':
                _, table_lines = block
                rows = parse_table(table_lines)
                add_table(doc, rows)
                doc.add_paragraph()  # spacing after table

    # Add references at the end
    if references_text:
        ref_blocks = parse_md_blocks(references_text)
        for block in ref_blocks:
            if block[0] == 'heading':
                p = doc.add_heading(block[2], level=block[1])
                for run in p.runs:
                    run.font.color.rgb = None
            elif block[0] == 'paragraph':
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(1.27)
                p.paragraph_format.first_line_indent = Cm(-1.27)
                add_rich_text(p, block[1])
                for run in p.runs:
                    run.font.size = Pt(10)

    out_path = ROOT / 'WIC_Methodology_and_Results.docx'
    doc.save(str(out_path))
    print(f'Saved to {out_path}')


if __name__ == '__main__':
    build_doc()
